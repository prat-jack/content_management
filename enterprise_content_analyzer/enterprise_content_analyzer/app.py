
import streamlit as st
import pandas as pd
import json
import io
import base64
from datetime import datetime
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from fpdf import FPDF
import PyPDF2
import os
from typing import Dict, List, Any

# Import the content analyzer with proper relative import
from .content_analyzer import ContentAnalyzer

class EnhancedContentAnalyzer(ContentAnalyzer):
    """Enhanced version of ContentAnalyzer with additional features"""
    
    def analyze_content_comprehensive(self, text: str) -> dict:
        """Enhanced analysis with more detailed metrics"""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": """You are an expert business document analyzer. Provide a comprehensive analysis in the following JSON format:
                        {
                            "summary": "Brief summary",
                            "sentiment": "positive/neutral/negative",
                            "confidence": 0.95,
                            "key_points": ["point1", "point2", "point3"],
                            "document_type": "email/report/memo/contract/etc",
                            "readability_score": 8.5,
                            "word_count": 500,
                            "complexity": "low/medium/high",
                            "recommendations": ["rec1", "rec2"],
                            "risk_factors": ["risk1", "risk2"],
                            "compliance_notes": "Any compliance considerations"
                        }"""
                    },
                    {
                        "role": "user",
                        "content": f"Analyze the following business document:\n\n{text}"
                    }
                ],
                max_tokens=800
            )
            
            try:
                analysis = json.loads(response.choices[0].message.content.strip())
                analysis['analysis_timestamp'] = datetime.now().isoformat()
                analysis['text_length'] = len(text)
                return analysis
            except json.JSONDecodeError:
                # Fallback to original format if JSON parsing fails
                return super().analyze_content(text)
                
        except Exception as e:
            return {"error": f"An error occurred: {e}"}

class ReportGenerator:
    """Handles export functionality for different formats"""
    
    @staticmethod
    def to_csv(analysis_data: Dict, content: str) -> str:
        """Convert analysis to CSV format"""
        df = pd.DataFrame([{
            'Timestamp': analysis_data.get('analysis_timestamp', datetime.now().isoformat()),
            'Content_Length': len(content),
            'Summary': analysis_data.get('summary', ''),
            'Sentiment': analysis_data.get('sentiment', ''),
            'Confidence': analysis_data.get('confidence', 0),
            'Document_Type': analysis_data.get('document_type', ''),
            'Complexity': analysis_data.get('complexity', ''),
            'Key_Points': '; '.join(analysis_data.get('key_points', [])),
            'Recommendations': '; '.join(analysis_data.get('recommendations', [])),
            'Risk_Factors': '; '.join(analysis_data.get('risk_factors', []))
        }])
        return df.to_csv(index=False)
    
    @staticmethod
    def to_json(analysis_data: Dict, content: str) -> str:
        """Convert analysis to JSON format"""
        export_data = {
            'content_sample': content[:200] + '...' if len(content) > 200 else content,
            'analysis': analysis_data,
            'export_timestamp': datetime.now().isoformat()
        }
        return json.dumps(export_data, indent=2)
    
    @staticmethod
    def to_pdf(analysis_data: Dict, content: str) -> bytes:
        """Convert analysis to PDF format"""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, 'Content Analysis Report', ln=True, align='C')
        
        pdf.set_font('Arial', '', 12)
        pdf.ln(10)
        
        # Add analysis details
        pdf.cell(0, 8, f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.cell(0, 8, f"Content Length: {len(content)} characters", ln=True)
        pdf.ln(5)
        
        # Summary
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 8, 'Summary:', ln=True)
        pdf.set_font('Arial', '', 10)
        summary_text = analysis_data.get('summary', 'No summary available')
        pdf.multi_cell(0, 6, summary_text)
        pdf.ln(5)
        
        # Key metrics
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 8, 'Key Metrics:', ln=True)
        pdf.set_font('Arial', '', 10)
        pdf.cell(0, 6, f"Sentiment: {analysis_data.get('sentiment', 'N/A')}", ln=True)
        pdf.cell(0, 6, f"Document Type: {analysis_data.get('document_type', 'N/A')}", ln=True)
        pdf.cell(0, 6, f"Complexity: {analysis_data.get('complexity', 'N/A')}", ln=True)
        
        return pdf.output(dest='S').encode('latin-1')
    
    @staticmethod
    def to_docx(analysis_data: Dict, content: str) -> bytes:
        """Convert analysis to DOCX format"""
        doc = Document()
        
        # Title
        title = doc.add_heading('Content Analysis Report', 0)
        title.alignment = 1  # Center alignment
        
        # Analysis details
        doc.add_heading('Analysis Overview', level=1)
        doc.add_paragraph(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Content Length: {len(content)} characters")
        
        # Summary
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(analysis_data.get('summary', 'No summary available'))
        
        # Key Points
        if analysis_data.get('key_points'):
            doc.add_heading('Key Points', level=1)
            for point in analysis_data['key_points']:
                doc.add_paragraph(point, style='List Bullet')
        
        # Metrics
        doc.add_heading('Analysis Metrics', level=1)
        metrics_table = doc.add_table(rows=1, cols=2)
        metrics_table.style = 'Table Grid'
        hdr_cells = metrics_table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        metrics = [
            ('Sentiment', analysis_data.get('sentiment', 'N/A')),
            ('Document Type', analysis_data.get('document_type', 'N/A')),
            ('Complexity', analysis_data.get('complexity', 'N/A')),
            ('Confidence Score', str(analysis_data.get('confidence', 'N/A')))
        ]
        
        for metric, value in metrics:
            row_cells = metrics_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
        
        # Save to bytes
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()

def extract_text_from_file(uploaded_file) -> str:
    """Extract text from various file formats"""
    try:
        file_type = uploaded_file.type
        
        if file_type == "text/plain":
            return str(uploaded_file.read(), "utf-8")
        
        elif file_type == "application/pdf":
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(uploaded_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        else:
            st.error(f"Unsupported file type: {file_type}")
            return ""
            
    except Exception as e:
        st.error(f"Error extracting text from file: {str(e)}")
        return ""

def create_visualizations(analysis_data: Dict):
    """Create visualizations for analysis results"""
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # Sentiment gauge
        sentiment_score = 0.5  # Default neutral
        if analysis_data.get('sentiment') == 'positive':
            sentiment_score = 0.8
        elif analysis_data.get('sentiment') == 'negative':
            sentiment_score = 0.2
        
        fig_gauge = go.Figure(go.Indicator(
            mode = "gauge+number",
            value = sentiment_score,
            domain = {'x': [0, 1], 'y': [0, 1]},
            title = {'text': "Sentiment Score"},
            gauge = {
                'axis': {'range': [None, 1]},
                'bar': {'color': "darkblue"},
                'steps': [
                    {'range': [0, 0.3], 'color': "lightgray"},
                    {'range': [0.3, 0.7], 'color': "gray"},
                    {'range': [0.7, 1], 'color': "lightgreen"}
                ],
                'threshold': {
                    'line': {'color': "red", 'width': 4},
                    'thickness': 0.75,
                    'value': 0.9
                }
            }
        ))
        fig_gauge.update_layout(height=300)
        st.plotly_chart(fig_gauge, use_container_width=True)
    
    with col2:
        # Confidence meter
        confidence = analysis_data.get('confidence', 0.5)
        fig_conf = go.Figure(go.Indicator(
            mode = "number+delta",
            value = confidence,
            title = {"text": "Confidence Level"},
            delta = {'reference': 0.7},
            number = {'suffix': "%", 'valueformat': ".1%"}
        ))
        fig_conf.update_layout(height=300)
        st.plotly_chart(fig_conf, use_container_width=True)
    
    with col3:
        # Document complexity
        complexity_levels = {'low': 1, 'medium': 2, 'high': 3}
        complexity = complexity_levels.get(analysis_data.get('complexity', 'medium'), 2)
        
        fig_complex = go.Figure(go.Bar(
            x=['Complexity'],
            y=[complexity],
            marker_color=['green' if complexity == 1 else 'orange' if complexity == 2 else 'red']
        ))
        fig_complex.update_layout(
            title="Document Complexity",
            yaxis=dict(range=[0, 3], tickvals=[1, 2, 3], ticktext=['Low', 'Medium', 'High']),
            height=300
        )
        st.plotly_chart(fig_complex, use_container_width=True)

def main():
    """Main Streamlit application"""
    
    # Page configuration
    st.set_page_config(
        page_title="Enterprise Content Analyzer",
        page_icon="📊",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS for better styling
    st.markdown("""
    <style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
    }
    .metric-card {
        background: #f8f9fa;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #667eea;
        margin: 0.5rem 0;
    }
    .export-section {
        background: #e3f2fd;
        padding: 1.5rem;
        border-radius: 10px;
        margin-top: 2rem;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🚀 Enterprise Content Analysis Platform</h1>
        <p>Advanced AI-powered document analysis with comprehensive export capabilities</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize analyzer
    analyzer = EnhancedContentAnalyzer()
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        analysis_mode = st.selectbox(
            "Analysis Mode",
            ["Comprehensive", "Quick", "Compliance Focus"],
            help="Choose the depth of analysis"
        )
        
        st.header("📤 Export Options")
        export_formats = st.multiselect(
            "Select Export Formats",
            ["CSV", "JSON", "PDF", "DOCX"],
            default=["JSON"],
            help="Choose formats for exporting results"
        )
        
        st.header("📊 Display Options")
        show_visualizations = st.checkbox("Show Visualizations", value=True)
        show_raw_data = st.checkbox("Show Raw Analysis Data", value=False)
        
        st.header("ℹ️ Information")
        st.info("This platform uses OpenAI's GPT-4 for content analysis. Ensure your API key is properly configured.")
    
    # Main content area
    tab1, tab2, tab3 = st.tabs(["📝 Text Analysis", "📁 File Upload", "📊 Analytics Dashboard"])
    
    with tab1:
        st.header("Text Content Analysis")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            content_input = st.text_area(
                "Enter your content here",
                height=300,
                placeholder="Paste your business document, email, report, or any text content here..."
            )
        
        with col2:
            st.markdown("### 📋 Analysis Features")
            st.markdown("""
            - **Sentiment Analysis**: Emotional tone detection
            - **Document Classification**: Automatic type identification
            - **Key Points Extraction**: Important information highlights
            - **Risk Assessment**: Potential issues identification
            - **Compliance Check**: Regulatory considerations
            - **Readability Score**: Content accessibility metrics
            """)
        
        if st.button("🔍 Analyze Content", type="primary", use_container_width=True):
            if content_input:
                with st.spinner("🧠 Analyzing content... Please wait"):
                    progress_bar = st.progress(0)
                    for i in range(100):
                        progress_bar.progress(i + 1)
                    
                    analysis_result = analyzer.analyze_content_comprehensive(content_input)
                    
                    if "error" in analysis_result:
                        st.error(f"❌ Analysis failed: {analysis_result['error']}")
                    else:
                        st.success("✅ Analysis completed successfully!")
                        
                        # Store results in session state for export
                        st.session_state['last_analysis'] = analysis_result
                        st.session_state['last_content'] = content_input
                        
                        # Display results
                        display_analysis_results(analysis_result, content_input, show_visualizations, show_raw_data, export_formats)
            else:
                st.warning("⚠️ Please enter some content to analyze.")
    
    with tab2:
        st.header("File Upload Analysis")
        
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['txt', 'pdf', 'docx'],
            help="Supported formats: TXT, PDF, DOCX"
        )
        
        if uploaded_file is not None:
            st.success(f"✅ File uploaded: {uploaded_file.name}")
            
            col1, col2 = st.columns([3, 1])
            
            with col1:
                if st.button("🔍 Analyze File", type="primary"):
                    with st.spinner("📄 Extracting text and analyzing..."):
                        extracted_text = extract_text_from_file(uploaded_file)
                        
                        if extracted_text:
                            st.info(f"📊 Extracted {len(extracted_text)} characters from file")
                            
                            analysis_result = analyzer.analyze_content_comprehensive(extracted_text)
                            
                            if "error" in analysis_result:
                                st.error(f"❌ Analysis failed: {analysis_result['error']}")
                            else:
                                st.success("✅ File analysis completed!")
                                
                                # Store results
                                st.session_state['last_analysis'] = analysis_result
                                st.session_state['last_content'] = extracted_text
                                
                                # Display results
                                display_analysis_results(analysis_result, extracted_text, show_visualizations, show_raw_data, export_formats)
                        else:
                            st.error("❌ Could not extract text from file")
            
            with col2:
                st.markdown("### 📄 File Info")
                st.markdown(f"**Name:** {uploaded_file.name}")
                st.markdown(f"**Size:** {uploaded_file.size} bytes")
                st.markdown(f"**Type:** {uploaded_file.type}")
    
    with tab3:
        st.header("Analytics Dashboard")
        
        if 'analysis_history' not in st.session_state:
            st.session_state['analysis_history'] = []
        
        if st.session_state.get('last_analysis'):
            # Add current analysis to history
            current_analysis = st.session_state['last_analysis'].copy()
            current_analysis['content_preview'] = st.session_state['last_content'][:100] + '...'
            
            if current_analysis not in st.session_state['analysis_history']:
                st.session_state['analysis_history'].append(current_analysis)
        
        if st.session_state['analysis_history']:
            # Display analytics
            df_history = pd.DataFrame(st.session_state['analysis_history'])
            
            col1, col2 = st.columns(2)
            
            with col1:
                # Sentiment distribution
                if 'sentiment' in df_history.columns:
                    fig_sent_dist = px.pie(
                        df_history, 
                        names='sentiment', 
                        title="Sentiment Distribution"
                    )
                    st.plotly_chart(fig_sent_dist, use_container_width=True)
            
            with col2:
                # Document types
                if 'document_type' in df_history.columns:
                    fig_doc_types = px.bar(
                        df_history['document_type'].value_counts().reset_index(),
                        x='index',
                        y='document_type',
                        title="Document Types Analyzed"
                    )
                    st.plotly_chart(fig_doc_types, use_container_width=True)
            
            # Analysis history table
            st.subheader("📚 Analysis History")
            st.dataframe(df_history[['analysis_timestamp', 'sentiment', 'document_type', 'complexity', 'content_preview']])
            
            if st.button("🗑️ Clear History"):
                st.session_state['analysis_history'] = []
                st.rerun()
        else:
            st.info("📊 No analysis history available yet. Analyze some content to see analytics.")

def display_analysis_results(analysis_result: Dict, content: str, show_visualizations: bool, show_raw_data: bool, export_formats: List[str]):
    """Display analysis results with formatting"""
    
    # Key metrics row
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown(f"""
        <div class="metric-card">
            <h4>📊 Sentiment</h4>
            <h2>{analysis_result.get('sentiment', 'N/A').title()}</h2>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown(f"""
        <div class="metric-card">
            <h4>📄 Document Type</h4>
            <h2>{analysis_result.get('document_type', 'N/A').title()}</h2>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown(f"""
        <div class="metric-card">
            <h4>🎯 Confidence</h4>
            <h2>{analysis_result.get('confidence', 0):.1%}</h2>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        st.markdown(f"""
        <div class="metric-card">
            <h4>⚡ Complexity</h4>
            <h2>{analysis_result.get('complexity', 'N/A').title()}</h2>
        </div>
        """, unsafe_allow_html=True)
    
    # Visualizations
    if show_visualizations:
        st.subheader("📊 Analysis Visualizations")
        create_visualizations(analysis_result)
    
    # Detailed results
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("📋 Analysis Summary")
        st.markdown(f"**Summary:** {analysis_result.get('summary', 'No summary available')}")
        
        if analysis_result.get('key_points'):
            st.markdown("**🔑 Key Points:**")
            for i, point in enumerate(analysis_result['key_points'], 1):
                st.markdown(f"{i}. {point}")
        
        if analysis_result.get('recommendations'):
            st.markdown("**💡 Recommendations:**")
            for rec in analysis_result['recommendations']:
                st.markdown(f"- {rec}")
    
    with col2:
        st.subheader("⚠️ Risk Assessment")
        if analysis_result.get('risk_factors'):
            for risk in analysis_result['risk_factors']:
                st.warning(f"⚠️ {risk}")
        else:
            st.success("✅ No significant risks identified")
        
        if analysis_result.get('compliance_notes'):
            st.info(f"📋 Compliance: {analysis_result['compliance_notes']}")
    
    # Raw data display
    if show_raw_data:
        st.subheader("🔍 Raw Analysis Data")
        st.json(analysis_result)
    
    # Export section
    if export_formats:
        st.markdown("""
        <div class="export-section">
            <h3>📤 Export Results</h3>
        </div>
        """, unsafe_allow_html=True)
        
        export_col1, export_col2, export_col3, export_col4 = st.columns(4)
        
        report_gen = ReportGenerator()
        
        if "CSV" in export_formats:
            with export_col1:
                csv_data = report_gen.to_csv(analysis_result, content)
                st.download_button(
                    label="📊 Download CSV",
                    data=csv_data,
                    file_name=f"analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv"
                )
        
        if "JSON" in export_formats:
            with export_col2:
                json_data = report_gen.to_json(analysis_result, content)
                st.download_button(
                    label="📄 Download JSON",
                    data=json_data,
                    file_name=f"analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json"
                )
        
        if "PDF" in export_formats:
            with export_col3:
                pdf_data = report_gen.to_pdf(analysis_result, content)
                st.download_button(
                    label="📑 Download PDF",
                    data=pdf_data,
                    file_name=f"analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf"
                )
        
        if "DOCX" in export_formats:
            with export_col4:
                docx_data = report_gen.to_docx(analysis_result, content)
                st.download_button(
                    label="📝 Download DOCX",
                    data=docx_data,
                    file_name=f"analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# Run the main app
if __name__ == "__main__":
    main()
